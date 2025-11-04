# -*- coding: utf-8 -*-
"""
Document generation service for care plans.
Supports PDF and Word document generation.
"""
from datetime import datetime
from typing import Dict, Any, Optional
from io import BytesIO
import os

from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.enums import TA_LEFT, TA_CENTER

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger


class DocumentGenerator:
    """ドキュメント生成サービス"""

    def __init__(self):
        """Initialize document generator."""
        self.setup_fonts()

    def setup_fonts(self):
        """日本語フォントのセットアップ"""
        try:
            # macOSの日本語フォントを登録
            font_path = "/System/Library/Fonts/ヒラギノ角ゴシック W3.ttc"
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont("JapaneseFont", font_path))
                self.japanese_font = "JapaneseFont"
            else:
                # フォールバック: Helvetica
                self.japanese_font = "Helvetica"
                logger.warning("Japanese font not found, using Helvetica")
        except Exception as e:
            logger.error(f"Error setting up fonts: {e}")
            self.japanese_font = "Helvetica"

    def generate_pdf(self, plan_data: Dict[str, Any], user_data: Dict[str, Any]) -> BytesIO:
        """
        Generate PDF care plan document.

        Args:
            plan_data: Plan information
            user_data: User information

        Returns:
            BytesIO: PDF document buffer
        """
        buffer = BytesIO()

        # ドキュメント作成
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
            leftMargin=15 * mm,
            rightMargin=15 * mm
        )

        # スタイル定義
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=self.japanese_font,
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=20
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=self.japanese_font,
            fontSize=14,
            spaceAfter=10
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontName=self.japanese_font,
            fontSize=10,
            leading=14
        )

        # コンテンツ構築
        story = []

        # タイトル
        title = Paragraph("サービス等利用計画書", title_style)
        story.append(title)
        story.append(Spacer(1, 10 * mm))

        # 利用者情報
        story.append(Paragraph("利用者情報", heading_style))
        user_info_data = [
            ["利用者名", user_data.get("name", "")],
            ["生年月日", f"{user_data.get('birth_date', '')} ({user_data.get('age', '')}歳)"],
            ["障害種別", user_data.get("disability_types", "")],
        ]
        user_table = Table(user_info_data, colWidths=[40 * mm, 120 * mm])
        user_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, -1), self.japanese_font, 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(user_table)
        story.append(Spacer(1, 10 * mm))

        # 長期目標
        story.append(Paragraph("長期目標", heading_style))
        for i, goal in enumerate(plan_data.get("long_term_goals", []), 1):
            goal_text = f"{i}. {goal.get('goal_text', '')}"
            story.append(Paragraph(goal_text, body_style))
            story.append(Paragraph(f"   評価期間: {goal.get('evaluation_period', '')}", body_style))
            story.append(Paragraph(f"   評価基準: {goal.get('evaluation_criteria', '')}", body_style))
            story.append(Spacer(1, 5 * mm))

        # 短期目標
        if plan_data.get("short_term_goals"):
            story.append(Paragraph("短期目標", heading_style))
            for i, goal in enumerate(plan_data.get("short_term_goals", []), 1):
                goal_text = f"{i}. {goal.get('goal_text', '')}"
                story.append(Paragraph(goal_text, body_style))
                story.append(Paragraph(f"   評価期間: {goal.get('evaluation_period', '')}", body_style))
                story.append(Spacer(1, 5 * mm))

        # サービス調整
        if plan_data.get("services"):
            story.append(Paragraph("サービス調整", heading_style))
            service_data = [["No", "事業所名", "サービス種別", "利用頻度", "開始予定日"]]
            for i, service in enumerate(plan_data.get("services", []), 1):
                service_data.append([
                    str(i),
                    service.get("facility_name", ""),
                    service.get("service_type", ""),
                    service.get("frequency", ""),
                    service.get("start_date", "")
                ])

            service_table = Table(service_data, colWidths=[10 * mm, 50 * mm, 40 * mm, 30 * mm, 30 * mm])
            service_table.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), self.japanese_font, 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
                ('LEFTPADDING', (0, 0), (-1, -1), 3),
                ('RIGHTPADDING', (0, 0), (-1, -1), 3),
            ]))
            story.append(service_table)

        # フッター情報
        story.append(Spacer(1, 15 * mm))
        footer_data = [
            ["作成日", datetime.now().strftime("%Y年%m月%d日")],
            ["計画ID", plan_data.get("plan_id", "")],
        ]
        footer_table = Table(footer_data, colWidths=[30 * mm, 130 * mm])
        footer_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, -1), self.japanese_font, 9),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(footer_table)

        # PDF生成
        doc.build(story)
        buffer.seek(0)
        return buffer

    def generate_word(self, plan_data: Dict[str, Any], user_data: Dict[str, Any]) -> BytesIO:
        """
        Generate Word care plan document.

        Args:
            plan_data: Plan information
            user_data: User information

        Returns:
            BytesIO: Word document buffer
        """
        buffer = BytesIO()
        doc = Document()

        # ドキュメント設定
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(11.69)  # A4 height
            section.page_width = Inches(8.27)    # A4 width
            section.top_margin = Inches(0.79)
            section.bottom_margin = Inches(0.79)
            section.left_margin = Inches(0.59)
            section.right_margin = Inches(0.59)

        # タイトル
        title = doc.add_heading("サービス等利用計画書", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 利用者情報
        doc.add_heading("利用者情報", level=2)
        user_table = doc.add_table(rows=3, cols=2)
        user_table.style = "Light Grid Accent 1"

        cells = user_table.rows[0].cells
        cells[0].text = "利用者名"
        cells[1].text = user_data.get("name", "")

        cells = user_table.rows[1].cells
        cells[0].text = "生年月日"
        cells[1].text = f"{user_data.get('birth_date', '')} ({user_data.get('age', '')}歳)"

        cells = user_table.rows[2].cells
        cells[0].text = "障害種別"
        cells[1].text = user_data.get("disability_types", "")

        doc.add_paragraph()

        # 長期目標
        doc.add_heading("長期目標", level=2)
        for i, goal in enumerate(plan_data.get("long_term_goals", []), 1):
            p = doc.add_paragraph(f"{i}. {goal.get('goal_text', '')}")
            p.style = "List Bullet"
            doc.add_paragraph(f"   評価期間: {goal.get('evaluation_period', '')}")
            doc.add_paragraph(f"   評価基準: {goal.get('evaluation_criteria', '')}")

        # 短期目標
        if plan_data.get("short_term_goals"):
            doc.add_heading("短期目標", level=2)
            for i, goal in enumerate(plan_data.get("short_term_goals", []), 1):
                p = doc.add_paragraph(f"{i}. {goal.get('goal_text', '')}")
                p.style = "List Bullet"
                doc.add_paragraph(f"   評価期間: {goal.get('evaluation_period', '')}")

        # サービス調整
        if plan_data.get("services"):
            doc.add_heading("サービス調整", level=2)
            service_table = doc.add_table(rows=len(plan_data.get("services", [])) + 1, cols=5)
            service_table.style = "Light Grid Accent 1"

            # ヘッダー
            headers = service_table.rows[0].cells
            headers[0].text = "No"
            headers[1].text = "事業所名"
            headers[2].text = "サービス種別"
            headers[3].text = "利用頻度"
            headers[4].text = "開始予定日"

            # データ
            for i, service in enumerate(plan_data.get("services", []), 1):
                row = service_table.rows[i].cells
                row[0].text = str(i)
                row[1].text = service.get("facility_name", "")
                row[2].text = service.get("service_type", "")
                row[3].text = service.get("frequency", "")
                row[4].text = service.get("start_date", "")

        # フッター
        doc.add_paragraph()
        doc.add_paragraph(f"作成日: {datetime.now().strftime('%Y年%m月%d日')}")
        doc.add_paragraph(f"計画ID: {plan_data.get('plan_id', '')}")

        # バッファに保存
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    def generate_monitoring_pdf(self, monitoring_data: Dict[str, Any], plan_data: Dict[str, Any], user_data: Dict[str, Any]) -> BytesIO:
        """
        Generate PDF monitoring record document.

        Args:
            monitoring_data: Monitoring record information
            plan_data: Plan information
            user_data: User information

        Returns:
            BytesIO: PDF document buffer
        """
        buffer = BytesIO()

        # ドキュメント作成
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
            leftMargin=15 * mm,
            rightMargin=15 * mm
        )

        # スタイル定義
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontName=self.japanese_font,
            fontSize=18,
            alignment=TA_CENTER,
            spaceAfter=20
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontName=self.japanese_font,
            fontSize=14,
            spaceAfter=10
        )
        body_style = ParagraphStyle(
            'CustomBody',
            parent=styles['Normal'],
            fontName=self.japanese_font,
            fontSize=10,
            leading=14
        )

        # コンテンツ構築
        story = []

        # タイトル
        title = Paragraph("モニタリング記録", title_style)
        story.append(title)
        story.append(Spacer(1, 10 * mm))

        # 基本情報
        story.append(Paragraph("基本情報", heading_style))
        basic_info_data = [
            ["利用者名", user_data.get("name", "")],
            ["記録日", str(monitoring_data.get("monitoring_date", ""))[:10]],
            ["種別", monitoring_data.get("monitoring_type", "")],
            ["ステータス", monitoring_data.get("status", "")],
        ]
        basic_table = Table(basic_info_data, colWidths=[40 * mm, 120 * mm])
        basic_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, -1), self.japanese_font, 10),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ('BACKGROUND', (0, 0), (0, -1), colors.lightgrey),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('LEFTPADDING', (0, 0), (-1, -1), 5),
            ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ]))
        story.append(basic_table)
        story.append(Spacer(1, 10 * mm))

        # 目標評価
        story.append(Paragraph("目標評価", heading_style))
        for i, goal_eval in enumerate(monitoring_data.get("goal_evaluations", []), 1):
            goal_type = "長期目標" if goal_eval.get("goal_type") == "long_term" else "短期目標"
            story.append(Paragraph(f"{goal_type} {i}", body_style))
            goal_eval_data = [
                ["達成率", f"{goal_eval.get('achievement_rate', 0)}%"],
                ["達成状況", goal_eval.get('achievement_status', '')],
                ["評価コメント", goal_eval.get('evaluation_comment', '')],
            ]
            if goal_eval.get('evidence'):
                goal_eval_data.append(["根拠", goal_eval.get('evidence')])
            if goal_eval.get('next_action'):
                goal_eval_data.append(["次のアクション", goal_eval.get('next_action')])

            goal_table = Table(goal_eval_data, colWidths=[30 * mm, 130 * mm])
            goal_table.setStyle(TableStyle([
                ('FONT', (0, 0), (-1, -1), self.japanese_font, 9),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                ('BACKGROUND', (0, 0), (0, -1), colors.lightblue),
                ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                ('LEFTPADDING', (0, 0), (-1, -1), 3),
                ('RIGHTPADDING', (0, 0), (-1, -1), 3),
            ]))
            story.append(goal_table)
            story.append(Spacer(1, 5 * mm))

        # サービス評価
        if monitoring_data.get("service_evaluations"):
            story.append(Paragraph("サービス評価", heading_style))
            for i, service_eval in enumerate(monitoring_data.get("service_evaluations", []), 1):
                story.append(Paragraph(f"サービス {i}: {service_eval.get('service_name', '')}", body_style))
                service_eval_data = [
                    ["出席率", f"{service_eval.get('attendance_rate', 0)}%"],
                    ["満足度", f"{service_eval.get('service_satisfaction', 0)}/5"],
                ]
                if service_eval.get('effectiveness'):
                    service_eval_data.append(["効果・変化", service_eval.get('effectiveness')])
                if service_eval.get('issues'):
                    service_eval_data.append(["課題", service_eval.get('issues')])

                service_table = Table(service_eval_data, colWidths=[30 * mm, 130 * mm])
                service_table.setStyle(TableStyle([
                    ('FONT', (0, 0), (-1, -1), self.japanese_font, 9),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
                    ('BACKGROUND', (0, 0), (0, -1), colors.lightgreen),
                    ('VALIGN', (0, 0), (-1, -1), 'TOP'),
                    ('LEFTPADDING', (0, 0), (-1, -1), 3),
                    ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                ]))
                story.append(service_table)
                story.append(Spacer(1, 5 * mm))

        # 総合評価
        story.append(Paragraph("総合評価", heading_style))
        story.append(Paragraph(monitoring_data.get("overall_summary", ""), body_style))
        story.append(Spacer(1, 5 * mm))

        if monitoring_data.get("strengths"):
            story.append(Paragraph("良かった点", body_style))
            for strength in monitoring_data.get("strengths", []):
                story.append(Paragraph(f"• {strength}", body_style))
            story.append(Spacer(1, 3 * mm))

        if monitoring_data.get("challenges"):
            story.append(Paragraph("課題", body_style))
            for challenge in monitoring_data.get("challenges", []):
                story.append(Paragraph(f"• {challenge}", body_style))
            story.append(Spacer(1, 3 * mm))

        if monitoring_data.get("family_feedback"):
            story.append(Paragraph("家族の意見", body_style))
            story.append(Paragraph(monitoring_data.get("family_feedback"), body_style))

        # フッター情報
        story.append(Spacer(1, 15 * mm))
        footer_data = [
            ["作成日", datetime.now().strftime("%Y年%m月%d日")],
            ["記録ID", monitoring_data.get("monitoring_id", "")],
        ]
        footer_table = Table(footer_data, colWidths=[30 * mm, 130 * mm])
        footer_table.setStyle(TableStyle([
            ('FONT', (0, 0), (-1, -1), self.japanese_font, 9),
            ('LEFTPADDING', (0, 0), (-1, -1), 0),
            ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ]))
        story.append(footer_table)

        # PDF生成
        doc.build(story)
        buffer.seek(0)
        return buffer

    def generate_monitoring_word(self, monitoring_data: Dict[str, Any], plan_data: Dict[str, Any], user_data: Dict[str, Any]) -> BytesIO:
        """
        Generate Word monitoring record document.

        Args:
            monitoring_data: Monitoring record information
            plan_data: Plan information
            user_data: User information

        Returns:
            BytesIO: Word document buffer
        """
        buffer = BytesIO()
        doc = Document()

        # ドキュメント設定
        sections = doc.sections
        for section in sections:
            section.page_height = Inches(11.69)  # A4 height
            section.page_width = Inches(8.27)    # A4 width
            section.top_margin = Inches(0.79)
            section.bottom_margin = Inches(0.79)
            section.left_margin = Inches(0.59)
            section.right_margin = Inches(0.59)

        # タイトル
        title = doc.add_heading("モニタリング記録", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # 基本情報
        doc.add_heading("基本情報", level=2)
        basic_table = doc.add_table(rows=4, cols=2)
        basic_table.style = "Light Grid Accent 1"

        cells = basic_table.rows[0].cells
        cells[0].text = "利用者名"
        cells[1].text = user_data.get("name", "")

        cells = basic_table.rows[1].cells
        cells[0].text = "記録日"
        cells[1].text = str(monitoring_data.get("monitoring_date", ""))[:10]

        cells = basic_table.rows[2].cells
        cells[0].text = "種別"
        cells[1].text = monitoring_data.get("monitoring_type", "")

        cells = basic_table.rows[3].cells
        cells[0].text = "ステータス"
        cells[1].text = monitoring_data.get("status", "")

        doc.add_paragraph()

        # 目標評価
        doc.add_heading("目標評価", level=2)
        for i, goal_eval in enumerate(monitoring_data.get("goal_evaluations", []), 1):
            goal_type = "長期目標" if goal_eval.get("goal_type") == "long_term" else "短期目標"
            p = doc.add_paragraph(f"{goal_type} {i}")
            p.style = "Heading 3"

            doc.add_paragraph(f"達成率: {goal_eval.get('achievement_rate', 0)}%")
            doc.add_paragraph(f"達成状況: {goal_eval.get('achievement_status', '')}")
            doc.add_paragraph(f"評価コメント: {goal_eval.get('evaluation_comment', '')}")
            if goal_eval.get('evidence'):
                doc.add_paragraph(f"根拠: {goal_eval.get('evidence')}")
            if goal_eval.get('next_action'):
                doc.add_paragraph(f"次のアクション: {goal_eval.get('next_action')}")
            doc.add_paragraph()

        # サービス評価
        if monitoring_data.get("service_evaluations"):
            doc.add_heading("サービス評価", level=2)
            for i, service_eval in enumerate(monitoring_data.get("service_evaluations", []), 1):
                p = doc.add_paragraph(f"サービス {i}: {service_eval.get('service_name', '')}")
                p.style = "Heading 3"

                doc.add_paragraph(f"出席率: {service_eval.get('attendance_rate', 0)}%")
                doc.add_paragraph(f"満足度: {service_eval.get('service_satisfaction', 0)}/5")
                if service_eval.get('effectiveness'):
                    doc.add_paragraph(f"効果・変化: {service_eval.get('effectiveness')}")
                if service_eval.get('issues'):
                    doc.add_paragraph(f"課題: {service_eval.get('issues')}")
                doc.add_paragraph()

        # 総合評価
        doc.add_heading("総合評価", level=2)
        doc.add_paragraph(monitoring_data.get("overall_summary", ""))

        if monitoring_data.get("strengths"):
            doc.add_paragraph("良かった点:")
            for strength in monitoring_data.get("strengths", []):
                doc.add_paragraph(f"• {strength}", style="List Bullet")

        if monitoring_data.get("challenges"):
            doc.add_paragraph("課題:")
            for challenge in monitoring_data.get("challenges", []):
                doc.add_paragraph(f"• {challenge}", style="List Bullet")

        if monitoring_data.get("family_feedback"):
            doc.add_paragraph("家族の意見:")
            doc.add_paragraph(monitoring_data.get("family_feedback"))

        # フッター
        doc.add_paragraph()
        doc.add_paragraph(f"作成日: {datetime.now().strftime('%Y年%m月%d日')}")
        doc.add_paragraph(f"記録ID: {monitoring_data.get('monitoring_id', '')}")

        # バッファに保存
        doc.save(buffer)
        buffer.seek(0)
        return buffer


def get_document_generator() -> DocumentGenerator:
    """Get document generator instance."""
    return DocumentGenerator()

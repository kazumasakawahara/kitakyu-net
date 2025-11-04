# -*- coding: utf-8 -*-
"""
Document generation service for plan forms.
"""
from datetime import datetime
from typing import Dict, Any, List, Optional
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from loguru import logger

from backend.services.user_service import get_user_service
from backend.services.assessment_service import get_assessment_service
from backend.services.goal_service import get_goal_service
from backend.services.service_need_service import get_service_need_service


class DocumentService:
    """Service for generating plan documents."""

    def __init__(self):
        """Initialize document service."""
        self.user_service = get_user_service()
        self.assessment_service = get_assessment_service()
        self.goal_service = get_goal_service()
        self.service_need_service = get_service_need_service()
        self.output_dir = Path("data/generated_documents")
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info("Document service initialized")

    def generate_plan_form1(
        self, plan_id: str, user_id: str, assessment_id: str, goal_ids: List[str], service_need_ids: List[str]
    ) -> str:
        """
        Generate Form 1 (サービス等利用計画書 様式1).

        Args:
            plan_id: Plan ID
            user_id: User ID
            assessment_id: Assessment ID
            goal_ids: List of goal IDs
            service_need_ids: List of service need IDs

        Returns:
            Path to generated document
        """
        logger.info(f"Generating Form 1 for plan: {plan_id}")

        # Gather data
        user = self.user_service.get_user(user_id)
        if not user:
            raise ValueError(f"User {user_id} not found")

        assessment = self.assessment_service.get_assessment(assessment_id)
        if not assessment:
            raise ValueError(f"Assessment {assessment_id} not found")

        goals = []
        for goal_id in goal_ids:
            goal = self.goal_service.get_goal(goal_id)
            if goal:
                goals.append(goal)

        service_needs = []
        for service_need_id in service_need_ids:
            service_need = self.service_need_service.get_service_need(service_need_id)
            if service_need:
                service_needs.append(service_need)

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading("サービス等利用計画書（様式1）", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Basic information
        doc.add_heading("利用者基本情報", level=2)
        self._add_table_row(doc, "利用者氏名", user["name"])
        self._add_table_row(doc, "生年月日", f"{user['birth_date']} ({user.get('age', '不明')}歳)")
        self._add_table_row(doc, "障害種別", user.get("disability_type", ""))
        self._add_table_row(doc, "障害支援区分", user.get("support_level", ""))
        self._add_table_row(doc, "居住形態", user.get("living_situation", ""))
        self._add_table_row(doc, "作成日", datetime.now().strftime("%Y年%m月%d日"))

        # Assessment results
        doc.add_heading("アセスメント結果", level=2)
        doc.add_paragraph(f"実施日: {assessment.get('interview_date', '')}")

        if assessment.get("preferences"):
            doc.add_paragraph("【本人の希望】")
            for pref in assessment["preferences"]:
                doc.add_paragraph(f"・{pref}", style='List Bullet')

        if assessment.get("analyzed_needs"):
            doc.add_paragraph("【分析されたニーズ】")
            for need in assessment["analyzed_needs"]:
                doc.add_paragraph(f"・{need}", style='List Bullet')

        if assessment.get("strengths"):
            doc.add_paragraph("【本人の強み】")
            for strength in assessment["strengths"]:
                doc.add_paragraph(f"・{strength}", style='List Bullet')

        if assessment.get("challenges"):
            doc.add_paragraph("【支援が必要な課題】")
            for challenge in assessment["challenges"]:
                doc.add_paragraph(f"・{challenge}", style='List Bullet')

        # Goals
        doc.add_heading("目標", level=2)

        long_term_goals = [g for g in goals if g.get("goal_type") == "長期目標"]
        short_term_goals = [g for g in goals if g.get("goal_type") == "短期目標"]

        if long_term_goals:
            doc.add_paragraph("【長期目標】")
            for goal in long_term_goals:
                doc.add_paragraph(f"・{goal['goal_text']}", style='List Bullet')
                if goal.get("goal_reason"):
                    doc.add_paragraph(f"  理由: {goal['goal_reason']}")
                if goal.get("evaluation_period"):
                    doc.add_paragraph(f"  評価期間: {goal['evaluation_period']}")

        if short_term_goals:
            doc.add_paragraph("【短期目標】")
            for goal in short_term_goals:
                doc.add_paragraph(f"・{goal['goal_text']}", style='List Bullet')
                if goal.get("goal_reason"):
                    doc.add_paragraph(f"  理由: {goal['goal_reason']}")
                if goal.get("evaluation_period"):
                    doc.add_paragraph(f"  評価期間: {goal['evaluation_period']}")

        # Service needs
        doc.add_heading("必要なサービス", level=2)
        for service_need in service_needs:
            doc.add_paragraph(f"【{service_need['service_type']}】", style='List Bullet')
            doc.add_paragraph(f"  頻度: {service_need['frequency']}")
            doc.add_paragraph(f"  優先度: {service_need['priority']}")
            doc.add_paragraph(f"  理由: {service_need['reason']}")
            if service_need.get("facility_name"):
                doc.add_paragraph(f"  事業所: {service_need['facility_name']}")

        # Save document
        output_path = self.output_dir / f"plan_form1_{plan_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(str(output_path))

        logger.success(f"Generated Form 1: {output_path}")
        return str(output_path)

    def generate_plan_form5(
        self, plan_id: str, user_id: str, service_need_ids: List[str]
    ) -> str:
        """
        Generate Form 5 (週間計画表 様式5).

        Args:
            plan_id: Plan ID
            user_id: User ID
            service_need_ids: List of service need IDs

        Returns:
            Path to generated document
        """
        logger.info(f"Generating Form 5 for plan: {plan_id}")

        # Gather data
        user = self.user_service.get_user(user_id)
        if not user:
            raise ValueError(f"User {user_id} not found")

        service_needs = []
        for service_need_id in service_need_ids:
            service_need = self.service_need_service.get_service_need(service_need_id)
            if service_need:
                service_needs.append(service_need)

        # Create document
        doc = Document()

        # Title
        title = doc.add_heading("週間計画表（様式5）", level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Basic information
        doc.add_heading("利用者情報", level=2)
        self._add_table_row(doc, "利用者氏名", user["name"])
        self._add_table_row(doc, "作成日", datetime.now().strftime("%Y年%m月%d日"))

        # Weekly schedule table
        doc.add_heading("週間スケジュール", level=2)

        # Create table: Days of week x Time slots
        table = doc.add_table(rows=5, cols=8)
        table.style = 'Light Grid Accent 1'

        # Header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "時間帯"
        days = ["月曜日", "火曜日", "水曜日", "木曜日", "金曜日", "土曜日", "日曜日"]
        for i, day in enumerate(days):
            header_cells[i + 1].text = day

        # Time slots
        time_slots = ["午前", "午後", "夕方", "夜間"]
        for i, time_slot in enumerate(time_slots):
            table.rows[i + 1].cells[0].text = time_slot

        # Service needs summary
        doc.add_heading("利用サービス一覧", level=2)
        for service_need in service_needs:
            doc.add_paragraph(f"【{service_need['service_type']}】", style='List Bullet')
            doc.add_paragraph(f"  頻度: {service_need['frequency']}")
            if service_need.get("duration_hours"):
                doc.add_paragraph(f"  時間: {service_need['duration_hours']}時間")
            if service_need.get("preferred_time"):
                doc.add_paragraph(f"  希望時間帯: {service_need['preferred_time']}")
            if service_need.get("facility_name"):
                doc.add_paragraph(f"  事業所: {service_need['facility_name']}")

        # Notes
        doc.add_paragraph()
        doc.add_paragraph("【備考】")
        doc.add_paragraph("※ 実際のスケジュールは利用者と事業所で調整してください。")
        doc.add_paragraph("※ 体調や状況により変更する場合があります。")

        # Save document
        output_path = self.output_dir / f"plan_form5_{plan_id}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        doc.save(str(output_path))

        logger.success(f"Generated Form 5: {output_path}")
        return str(output_path)

    def _add_table_row(self, doc: Document, label: str, value: str):
        """Add a labeled table row to document."""
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        cells = table.rows[0].cells
        cells[0].text = label
        cells[1].text = str(value)
        cells[0].width = Inches(2.0)
        cells[1].width = Inches(4.0)


# Global service instance
_document_service = None


def get_document_service() -> DocumentService:
    """Get or create document service instance."""
    global _document_service
    if _document_service is None:
        _document_service = DocumentService()
    return _document_service
